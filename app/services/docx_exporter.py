"""Word (.docx) export service using python-docx."""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

logger = logging.getLogger(__name__)


def export_to_docx(
    question: str,
    confidence_grade: str | None,
    execution_plan: str | None,
    created_at: str | None,
    findings: list[dict],
    narrative: dict,
    output_path: str | Path,
) -> Path:
    """Build a structured Word document from analysis results.

    Uses plain parameters instead of the PipelineRun model so the exporter
    stays decoupled from the ORM layer.

    Args:
        question: The pipeline question used as the document title.
        confidence_grade: Letter grade (A-F) or None.
        execution_plan: Name of the execution plan or None.
        created_at: ISO-formatted creation timestamp or None.
        findings: List of finding dicts, each with headline, impact, and detail.
        narrative: Dict with executive_summary and detailed_findings keys.
        output_path: Destination file path for the generated .docx.

    Returns:
        The resolved output path.

    Raises:
        RuntimeError: If document generation fails.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        doc = Document()

        # 1. Title — the pipeline question
        doc.add_heading(question, level=0)

        # 2. Metadata paragraph
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.LEFT

        grade_run = meta.add_run(f"Confidence: {confidence_grade or 'N/A'}")
        grade_run.font.size = Pt(10)
        grade_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        meta.add_run("  |  ")

        plan_run = meta.add_run(f"Plan: {execution_plan or 'deep_dive'}")
        plan_run.font.size = Pt(10)
        plan_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        meta.add_run("  |  ")

        date_run = meta.add_run(f"Date: {created_at or 'N/A'}")
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        # 3. Executive Summary
        doc.add_heading("Executive Summary", level=1)
        exec_summary = narrative.get("executive_summary", "") if isinstance(narrative, dict) else ""
        doc.add_paragraph(exec_summary)

        # 4. Findings
        doc.add_heading("Findings", level=1)
        for finding in findings:
            headline = finding.get("headline", "Finding")
            impact = finding.get("impact", "medium")
            detail = finding.get("detail", "")

            doc.add_heading(headline, level=2)

            impact_para = doc.add_paragraph()
            impact_label = impact_para.add_run("Impact: ")
            impact_label.bold = True
            impact_para.add_run(impact)

            doc.add_paragraph(detail)

        # 5. Detailed Analysis
        doc.add_heading("Detailed Analysis", level=1)
        detailed = narrative.get("detailed_findings", "") if isinstance(narrative, dict) else ""
        doc.add_paragraph(detailed)

        doc.save(str(output_path))

    except Exception as exc:
        logger.error("DOCX generation failed: %s", exc)
        raise RuntimeError(f"DOCX generation failed: {exc}") from exc

    return output_path
